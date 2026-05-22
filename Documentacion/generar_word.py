"""
Genera DOCUMENTACION_CRONOGRULLA.docx a partir del .md
"""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import os as _os
_DIR    = _os.path.dirname(_os.path.abspath(__file__))
MD_FILE = _os.path.join(_DIR, "DOCUMENTACION_CRONOGRULLA.md")
OUT_FILE = _os.path.join(_DIR, "DOCUMENTACION_CRONOGRULLA.docx")

# ── Colores de la paleta CronoGrulla ─────────────────────────────────────────
C_AZUL      = RGBColor(0x2C, 0x3E, 0x50)   # Sidebar oscuro
C_AZUL_MED  = RGBColor(0x34, 0x98, 0xDB)   # Azul acento
C_VERDE     = RGBColor(0x27, 0xAE, 0x60)   # Verde éxito
C_BLANCO    = RGBColor(0xFF, 0xFF, 0xFF)
C_GRIS_CLR  = RGBColor(0xF4, 0xF6, 0xF8)   # Fondo filas impares tabla
C_TEXTO     = RGBColor(0x2C, 0x3E, 0x50)   # Texto principal


# ── Helpers XML ───────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    """Rellena el fondo de una celda con color hexadecimal."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge, attrs in kwargs.items():
        edge_el = OxmlElement(f'w:{edge}')
        for k, v in attrs.items():
            edge_el.set(qn(f'w:{k}'), v)
        tcBorders.append(edge_el)
    tcPr.append(tcBorders)


def add_page_border(doc):
    """Añade borde fino a todas las páginas."""
    for section in doc.sections:
        sectPr = section._sectPr
        pgBorders = OxmlElement('w:pgBorders')
        pgBorders.set(qn('w:offsetFrom'), 'page')
        for side in ('top', 'left', 'bottom', 'right'):
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'),   'single')
            border.set(qn('w:sz'),    '4')
            border.set(qn('w:space'), '24')
            border.set(qn('w:color'), '2C3E50')
            pgBorders.append(border)
        sectPr.append(pgBorders)


# ── Estilos base ──────────────────────────────────────────────────────────────
def apply_base_styles(doc: Document):
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.font.color.rgb = C_TEXTO

    for level, sz, bold, color in [
        ('Heading 1', 18, True,  C_AZUL),
        ('Heading 2', 14, True,  C_AZUL_MED),
        ('Heading 3', 12, True,  C_VERDE),
        ('Heading 4', 11, True,  C_TEXTO),
    ]:
        h = doc.styles[level]
        h.font.name  = 'Calibri'
        h.font.size  = Pt(sz)
        h.font.bold  = bold
        h.font.color.rgb = color
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after  = Pt(4)


# ── Portada ───────────────────────────────────────────────────────────────────
def add_cover(doc: Document):
    # Bloque de color superior
    cover_p = doc.add_paragraph()
    cover_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover_p.add_run()
    run.add_break()

    # Emoji + Título principal
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title_p.add_run("📋  DOCUMENTACIÓN TÉCNICA")
    r.font.size  = Pt(28)
    r.font.bold  = True
    r.font.color.rgb = C_AZUL

    # Subtítulo
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub_p.add_run("CronoGrulla — Sistema de Ingeniería de Métodos y Estudio de Tiempos")
    rs.font.size  = Pt(16)
    rs.font.color.rgb = C_AZUL_MED

    # Meta info
    doc.add_paragraph()
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rm = meta_p.add_run("Versión 1.0  |  Abril 2026  |  Universidad Católica")
    rm.font.size  = Pt(12)
    rm.font.color.rgb = C_GRIS_CLR
    rm.font.italic = True

    # Línea separadora
    doc.add_paragraph()
    sep = doc.add_paragraph("─" * 65)
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep.runs[0].font.color.rgb = C_AZUL_MED

    doc.add_page_break()


# ── Tabla Markdown → Word ─────────────────────────────────────────────────────
def parse_md_table(lines: list[str], doc: Document):
    rows = []
    for line in lines:
        if line.startswith('|'):
            cells = [c.strip() for c in line.strip().strip('|').split('|')]
            rows.append(cells)

    # Eliminar fila separadora (---|---)
    rows = [r for r in rows if not all(re.match(r'^[-:]+$', c) for c in r)]
    if not rows:
        return

    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'

    col_widths = [Inches(6.4 / num_cols)] * num_cols

    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.width = col_widths[j]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            # Limpiar texto de markdown inline
            clean = re.sub(r'\*\*(.*?)\*\*', r'\1', cell_text)
            clean = re.sub(r'`(.*?)`', r'\1', clean)

            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(clean)
            run.font.size = Pt(10)

            if i == 0:
                # Encabezado: fondo azul oscuro, texto blanco, negrita
                set_cell_bg(cell, '2C3E50')
                run.font.bold  = True
                run.font.color.rgb = C_BLANCO
            elif i % 2 == 0:
                set_cell_bg(cell, 'EBF5FB')
                run.font.color.rgb = C_TEXTO
            else:
                set_cell_bg(cell, 'FDFEFE')
                run.font.color.rgb = C_TEXTO

    doc.add_paragraph()


# ── Bloque de código ──────────────────────────────────────────────────────────
def add_code_block(doc: Document, code_lines: list[str]):
    for line in code_lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        run = p.add_run(line if line else " ")
        run.font.name  = 'Courier New'
        run.font.size  = Pt(9)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        # Fondo gris suave vía shading en párrafo
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  'F2F3F4')
        pPr.append(shd)
    doc.add_paragraph()


# ── Texto con inline markdown (negrita, código) ───────────────────────────────
def add_paragraph_with_inline(doc: Document, text: str, bold=False, size=11, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    parts = re.split(r'(\*\*.*?\*\*|`.*?`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2])
            r.bold = True
        elif part.startswith('`') and part.endswith('`'):
            r = p.add_run(part[1:-1])
            r.font.name = 'Courier New'
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        else:
            r = p.add_run(part)
        r.font.size = Pt(size)
        if bold:
            r.bold = True
        if color:
            r.font.color.rgb = color
    return p


# ── Procesador principal ──────────────────────────────────────────────────────
def process_md(doc: Document, md_path: str):
    with open(md_path, encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\n')

        # ── Encabezados ──────────────────────────────────────────────────────
        if line.startswith('#### '):
            p = doc.add_heading(line[5:].strip(), level=4)
        elif line.startswith('### '):
            p = doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('## '):
            p = doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('# '):
            # Ya tenemos la portada; ignorar el H1 duplicado
            pass

        # ── Separador horizontal ──────────────────────────────────────────────
        elif line.strip() == '---':
            sep = doc.add_paragraph("─" * 65)
            sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sep.runs[0].font.color.rgb = C_AZUL_MED
            sep.paragraph_format.space_before = Pt(6)
            sep.paragraph_format.space_after  = Pt(6)

        # ── Bloque de código ──────────────────────────────────────────────────
        elif line.startswith('```'):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].rstrip('\n').startswith('```'):
                code_lines.append(lines[i].rstrip('\n'))
                i += 1
            add_code_block(doc, code_lines)

        # ── Tabla ─────────────────────────────────────────────────────────────
        elif line.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].rstrip('\n').startswith('|'):
                table_lines.append(lines[i].rstrip('\n'))
                i += 1
            parse_md_table(table_lines, doc)
            continue  # i ya avanzó

        # ── Lista ─────────────────────────────────────────────────────────────
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Cm(1)
            text = line[2:]
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'`(.*?)`', r'\1', text)
            run = p.add_run(text)
            run.font.size = Pt(11)

        # ── Texto normal ──────────────────────────────────────────────────────
        elif line.strip():
            # Línea de metadatos de portada (Version, Fecha)
            if line.startswith('Versión:'):
                pass  # ya está en la portada
            else:
                add_paragraph_with_inline(doc, line)

        # ── Línea vacía ───────────────────────────────────────────────────────
        else:
            pass  # No agregar párrafos vacíos innecesarios

        i += 1


# ── Main ──────────────────────────────────────────────────────────────────────
def main():
    import os
    doc = Document()

    # Márgenes de página
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    apply_base_styles(doc)
    add_page_border(doc)
    add_cover(doc)
    process_md(doc, MD_FILE)

    doc.save(OUT_FILE)
    print(f"✅ Documento generado: {os.path.abspath(OUT_FILE)}")


if __name__ == '__main__':
    main()
