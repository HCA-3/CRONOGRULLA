import docx
import os

file_path = r"c:\Users\dsant\Desktop\U CATOLICA\ING INDUSTRIAL\SEMESTRES INDUSTRIAL\segundo semestre\ingenieria de metodos\Corte 1\ingenieria de metodos\Codigo.docx"

if os.path.exists(file_path):
    try:
        doc = docx.Document(file_path)
        print("Codigo.docx successfully opened!")
        print("Number of paragraphs:", len(doc.paragraphs))
        print("Number of tables:", len(doc.tables))
        print("\n--- First 20 paragraphs ---")
        for i, para in enumerate(doc.paragraphs[:20]):
            print(f"{i}: {para.text[:120]}")
    except Exception as e:
        print("Error reading Codigo.docx:", e)
else:
    print("Codigo.docx does not exist!")
