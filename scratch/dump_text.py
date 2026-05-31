"""
Script para extraer TODO el texto del documento IEEE y analizar
los errores de ortografia y codificacion.
"""
import docx
import sys
import io

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

p = r"C:\Users\dsant\Desktop\esta\Articulo_CronoGrulla_IEEE.docx"
doc = docx.Document(p)

print("=== TEXTO COMPLETO DE CADA PARRAFO ===\n")
for i, para in enumerate(doc.paragraphs):
    text = para.text
    if text.strip():
        # Mostrar caracteres problematicos
        problems = []
        for j, c in enumerate(text):
            code = ord(c)
            if code > 127 and code < 256:
                problems.append(f"  pos {j}: chr({code}) = '{c}' -> contexto: ...{text[max(0,j-5):j+6]}...")
            elif code > 8000:
                problems.append(f"  pos {j}: chr({code}) = '{c}' -> contexto: ...{text[max(0,j-5):j+6]}...")
        
        print(f"--- PARRAFO [{i}] (style: {para.style.name}) ---")
        print(text)
        if problems:
            print("CARACTERES ESPECIALES:")
            for prob in problems:
                print(prob)
        print()

print("\n=== TEXTO DE TABLAS ===\n")
for t, table in enumerate(doc.tables):
    print(f"--- TABLA {t} ---")
    for r, row in enumerate(table.rows):
        for c, cell in enumerate(row.cells):
            text = cell.text.strip()
            if text:
                print(f"  [{r},{c}] {text}")
    print()
