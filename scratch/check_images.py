import docx

file_path = r"c:\Users\dsant\Desktop\U CATOLICA\ING INDUSTRIAL\SEMESTRES INDUSTRIAL\segundo semestre\ingenieria de metodos\Corte 1\ingenieria de metodos\Informe Final Métodos.docx"
doc = docx.Document(file_path)

print("Listing all paragraphs that have inline shapes or drawings:")
for idx, p in enumerate(doc.paragraphs):
    # Check runs for inline shapes
    has_image = False
    for r in p.runs:
        if r.element.xpath('.//w:drawing') or r.element.xpath('.//w:pict') or r.element.xpath('.//w:object'):
            has_image = True
            break
    if has_image:
        print(f"Paragraph [{idx}] has image! Text around it: {p.text.strip() or '(empty)'}")

print("\nListing all tables and checking if there are images in cells:")
for idx, tbl in enumerate(doc.tables):
    for r_idx, r in enumerate(tbl.rows):
        for c_idx, cell in enumerate(r.cells):
            for p in cell.paragraphs:
                has_image = False
                for run in p.runs:
                    if run.element.xpath('.//w:drawing') or run.element.xpath('.//w:pict') or run.element.xpath('.//w:object'):
                        has_image = True
                        break
                if has_image:
                    print(f"Table [{idx}] Row [{r_idx}] Col [{c_idx}] has image! Text: {p.text.strip() or '(empty)'}")
