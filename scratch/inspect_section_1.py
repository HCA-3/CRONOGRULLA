import docx

file_path = r"c:\Users\dsant\Desktop\U CATOLICA\ING INDUSTRIAL\SEMESTRES INDUSTRIAL\segundo semestre\ingenieria de metodos\Corte 1\ingenieria de metodos\Informe Final Métodos.docx"
doc = docx.Document(file_path)

print("Paragraphs between 38 and 48:")
for i in range(38, 48):
    if i < len(doc.paragraphs):
        print(f"{i}: {doc.paragraphs[i].text}")
