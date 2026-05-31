import docx

file_path = r"c:\Users\dsant\Desktop\U CATOLICA\ING INDUSTRIAL\SEMESTRES INDUSTRIAL\segundo semestre\ingenieria de metodos\Corte 1\ingenieria de metodos\Informe Final Métodos.docx"
doc = docx.Document(file_path)

print("Total tables in document:", len(doc.tables))
for idx, table in enumerate(doc.tables):
    print(f"\n--- Table {idx} ---")
    print(f"Rows: {len(table.rows)}, Cols: {len(table.columns)}")
    # Print the first row (headers)
    headers = [cell.text.strip().replace("\n", " ") for cell in table.rows[0].cells]
    print("Header:", headers)
    # Print second row
    if len(table.rows) > 1:
        row1 = [cell.text.strip().replace("\n", " ") for cell in table.rows[1].cells]
        print("Row 1 :", row1)
