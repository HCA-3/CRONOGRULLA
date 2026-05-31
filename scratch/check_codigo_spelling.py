import docx
import os
import re

file_path = r"c:\Users\dsant\Desktop\U CATOLICA\ING INDUSTRIAL\SEMESTRES INDUSTRIAL\segundo semestre\ingenieria de metodos\Corte 1\ingenieria de metodos\Codigo.docx"

if not os.path.exists(file_path):
    print("File not found")
    exit()

doc = docx.Document(file_path)
print(f"Total Paragraphs: {len(doc.paragraphs)}")
print(f"Total Tables: {len(doc.tables)}")

# Look for common spelling errors or corrupted characters
count = 0
for idx, p in enumerate(doc.paragraphs):
    text = p.text
    if '' in text or '?' in text or any(w in text.lower() for w in ['metodologia', 'analisis', 'ptimo']):
        print(f"P[{idx}]: {text[:120]}...")
        count += 1
        if count > 20:
            print("Truncated list...")
            break
            
print(f"Found {count} matches.")
