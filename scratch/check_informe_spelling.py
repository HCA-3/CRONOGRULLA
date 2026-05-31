import docx
import os
import re

file_path = r"c:\Users\dsant\Desktop\U CATOLICA\ING INDUSTRIAL\SEMESTRES INDUSTRIAL\segundo semestre\ingenieria de metodos\Corte 1\ingenieria de metodos\Informe Final Métodos.docx"

if not os.path.exists(file_path):
    print("File not found")
    exit()

doc = docx.Document(file_path)
print(f"Total Paragraphs: {len(doc.paragraphs)}")
print(f"Total Tables: {len(doc.tables)}")

# Look for common spelling errors
typos = [
    r'\bmetodologia\b', r'\bmetodologias\b', r'\bproduccion\b', r'\bcalificacion\b',
    r'\boperario\b', r'\boperarios\b', r'\bestandar\b', r'\bdefinicion\b',
    r'\btecnica\b', r'\btecnicas\b', r'\bmedicion\b', r'\bmediciones\b',
    r'\bclasificacion\b', r'\bclasificaciones\b', r'\banalisis\b', r'\bgrafico\b',
    r'\bgraficos\b', r'\bcamara\b', r'\bcamaras\b', r'\bduracion\b', r'\blinea\b',
    r'\blineas\b', r'\bestudio\b', r'\bergonomia\b', r'\bergonomico\b',
    r'\bergonomica\b', r'\bergonomicos\b', r'\bergonomicas\b', r'\bptimo\b',
    r'\bptima\b', r'\bptimos\b', r'\bptimas\b'
]

print("\n--- Checking paragraphs for unaccented keywords ---")
count = 0
for idx, p in enumerate(doc.paragraphs):
    text = p.text
    for pattern in typos:
        if re.search(pattern, text, re.IGNORECASE):
            print(f"P[{idx}]: {text[:120]}...")
            count += 1
            break
            
print(f"Found {count} matches in paragraphs.")

print("\n--- Checking tables for unaccented keywords ---")
tbl_count = 0
for t_idx, table in enumerate(doc.tables):
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            text = cell.text
            for pattern in typos:
                if re.search(pattern, text, re.IGNORECASE):
                    print(f"Table[{t_idx}] Cell[{r_idx},{c_idx}]: {text[:80]}...")
                    tbl_count += 1
                    break
print(f"Found {tbl_count} matches in tables.")
