import docx

file_path = r"c:\Users\dsant\Desktop\U CATOLICA\ING INDUSTRIAL\SEMESTRES INDUSTRIAL\segundo semestre\ingenieria de metodos\Corte 1\ingenieria de metodos\Informe Final Métodos.docx"
doc = docx.Document(file_path)

print("Searching paragraphs...")
for idx, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    if any(q in txt.lower() for q in ["1.2", "micro-movimientos", "micromovimientos", "therblig"]):
        print(f"P {idx}: {txt}")

print("\nSearching table cells...")
for t_idx, table in enumerate(doc.tables):
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            txt = cell.text.strip()
            if any(q in txt.lower() for q in ["1.2", "micro-movimientos", "micromovimientos", "therblig"]):
                print(f"Table {t_idx}, Row {r_idx}, Col {c_idx}: {txt[:100]}...")
