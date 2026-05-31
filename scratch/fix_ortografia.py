"""
Script COMPLETO para corregir la ortografia del documento IEEE CronoGrulla.
Guarda en archivo temporal y luego reemplaza.
"""

import os
import sys
import io
import re
import shutil
from docx import Document

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

DOC_PATH = r"C:\Users\dsant\Desktop\esta\Articulo_CronoGrulla_IEEE.docx"
TEMP_PATH = r"C:\Users\dsant\Desktop\esta\Articulo_CronoGrulla_IEEE_CORREGIDO.docx"

doc = Document(DOC_PATH)
print(f"Documento abierto: {len(doc.paragraphs)} parrafos, {len(doc.tables)} tablas")

# ============================================================
# MAPA de caracteres corruptos
# ============================================================
CHAR_MAP = {
    chr(225): 'á', chr(233): 'é', chr(237): 'í', chr(243): 'ó', chr(250): 'ú',
    chr(241): 'ñ', chr(193): 'Á', chr(201): 'É', chr(205): 'Í', chr(211): 'Ó',
    chr(218): 'Ú', chr(209): 'Ñ', chr(252): 'ü', chr(220): 'Ü',
    chr(173): '', chr(150): '–', chr(151): '—',
}

def fix_chars(text):
    if not text:
        return text
    return "".join(CHAR_MAP.get(c, c) for c in text)

# ============================================================
# Patrones de doble-corrupcion
# ============================================================
DOUBLE_PATTERNS = [
    ("estááudio", "estudio"), ("estáá udio", "estudio"),
    ("estááe ", "este "), ("estáá a ", "esta "),
    ("estááa ", "esta "), ("estááo ", "esto "),
    ("estááas ", "estas "), ("estááos ", "estos "),
    ("estááaci", "estaci"), ("estáárecha", "estrecha"),
    ("estááima", "estima"), ("estááablecido", "establecido"),
    ("Puestááo", "Puesto"), ("puestááo", "puesto"),
    ("puestááas", "puestas"), ("propuestááas", "propuestas"),
    ("demuestááran", "demuestran"),
    ("estáá?ndar", "estándar"),
    # Doble "á" genérico - catch residuales
]

def fix_double(text):
    if not text:
        return text
    for old, new in DOUBLE_PATTERNS:
        text = text.replace(old, new)
    # Patron generico: "estáá" seguido de letras -> "est" + letras
    text = re.sub(r'estáá(\w)', r'est\1', text)
    return text

# ============================================================
# Correcciones ortográficas (español colombiano)
# ============================================================
SPELLING = [
    # --- Palabras terminadas en -ción ---
    ("Introduccion", "Introducción"), ("introduccion", "introducción"),
    ("Deteccion", "Detección"), ("deteccion", "detección"),
    ("Clasificacion", "Clasificación"), ("clasificacion", "clasificación"),
    ("Verificacion", "Verificación"), ("verificacion", "verificación"),
    ("Regresion", "Regresión"), ("regresion", "regresión"),
    ("Prediccion", "Predicción"), ("prediccion", "predicción"),
    ("Evaluacion", "Evaluación"), ("evaluacion", "evaluación"),
    ("Implementacion", "Implementación"), ("implementacion", "implementación"),
    ("Operacion", "Operación"), ("operacion", "operación"),
    ("Solucion", "Solución"), ("solucion", "solución"),
    ("Accion", "Acción"), ("accion", "acción"),
    ("acciones", "acciones"),
    ("Precaucion", "Precaución"), ("precaucion", "precaución"),
    ("Proyeccion", "Proyección"), ("proyeccion", "proyección"),
    ("Alimentacion", "Alimentación"), ("alimentacion", "alimentación"),
    ("Seccion", "Sección"), ("seccion", "sección"),
    ("Produccion", "Producción"), ("produccion", "producción"),
    ("Estimacion", "Estimación"), ("estimacion", "estimación"),
    ("Posicion", "Posición"), ("posicion", "posición"),
    ("Informacion", "Información"), ("informacion", "información"),
    ("Investigacion", "Investigación"), ("investigacion", "investigación"),
    ("Correlacion", "Correlación"), ("correlacion", "correlación"),
    ("Interaccion", "Interacción"), ("interaccion", "interacción"),
    ("Integracion", "Integración"), ("integracion", "integración"),
    ("Configuracion", "Configuración"), ("configuracion", "configuración"),
    ("Comunicacion", "Comunicación"), ("comunicacion", "comunicación"),
    ("Aplicacion", "Aplicación"), ("aplicacion", "aplicación"),
    ("Observacion", "Observación"), ("observacion", "observación"),
    ("manipulacion", "manipulación"), ("Manipulacion", "Manipulación"),
    ("activacion", "activación"), ("Activacion", "Activación"),
    ("visualizacion", "visualización"), ("Visualizacion", "Visualización"),
    ("medicion", "medición"), ("Medicion", "Medición"),
    ("composicional", "composicional"),
    ("modulacion", "modulación"),
    
    # --- Palabras terminadas en -ía ---
    ("Ingenieria", "Ingeniería"), ("ingenieria", "ingeniería"),
    ("Metodologia", "Metodología"), ("metodologia", "metodología"),
    ("Ergonomia", "Ergonomía"), ("ergonomia", "ergonomía"),
    ("cronometria", "cronometría"), ("Cronometria", "Cronometría"),
    ("Tecnologia", "Tecnología"), ("tecnologia", "tecnología"),
    ("categoria", "categoría"), ("Categoria", "Categoría"),
    ("geometria", "geometría"),
    
    # --- Adjetivos con tilde ---
    ("Analisis", "Análisis"), ("analisis", "análisis"),
    ("Ergonomico", "Ergonómico"), ("ergonomico", "ergonómico"),
    ("ergonomica", "ergonómica"), ("ergonomicas", "ergonómicas"),
    ("ergonomicos", "ergonómicos"),
    ("metodologica", "metodológica"), ("Metodologica", "Metodológica"),
    ("Diseno", "Diseño"), ("diseno", "diseño"),
    ("optica", "óptica"), ("optico", "óptico"), ("Optica", "Óptica"),
    ("angulo ", "ángulo "), ("angulos", "ángulos"), ("Angulo", "Ángulo"),
    ("atomos", "átomos"),
    ("ciberfisico", "ciberfísico"), ("ciberfisicos", "ciberfísicos"),
    ("ciberfisica", "ciberfísica"),
    ("esqueletico", "esquelético"), ("esqueletica", "esquelética"),
    ("biomecanico", "biomecánico"), ("biomecanica", "biomecánica"),
    ("problematica", "problemática"),
    ("tecnologico", "tecnológico"), ("tecnologica", "tecnológica"),
    ("cinematica", "cinemática"), ("cinematicas", "cinemáticas"),
    ("cinematico", "cinemático"),
    ("empirica", "empírica"), ("empirico", "empírico"),
    ("grafica", "gráfica"), ("graficas", "gráficas"),
    ("Grafica", "Gráfica"), ("Graficas", "Gráficas"),
    ("geometrica", "geométrica"), ("geometrico", "geométrico"),
    ("analitico", "analítico"), ("analitica", "analítica"),
    ("Analitico", "Analítico"),
    ("automatico", "automático"), ("automatica", "automática"),
    ("dinamico", "dinámico"), ("dinamica", "dinámica"),
    ("clasica", "clásica"), ("clasico", "clásico"),
    ("critico", "crítico"), ("critica", "crítica"),
    ("Catolica", "Católica"), ("catolica", "católica"),
    ("sistematico", "sistemático"), ("sistematica", "sistemática"),
    ("jerarquica", "jerárquica"), ("jerarquico", "jerárquico"),
    ("matematico", "matemático"), ("matematica", "matemática"),
    
    # --- Sustantivos con tilde ---
    ("Articulo", "Artículo"), ("articulo", "artículo"),
    ("metodos", "métodos"), ("Metodos", "Métodos"),
    ("metodo", "método"), ("Metodo", "Método"),
    ("indice", "índice"), ("Indice", "Índice"),
    ("calculo", "cálculo"), ("Calculo", "Cálculo"),
    ("maquina", "máquina"), ("Maquina", "Máquina"),
    ("modulo", "módulo"), ("Modulo", "Módulo"),
    ("modulos", "módulos"), ("Modulos", "Módulos"),
    ("sesion", "sesión"),
    ("estandar", "estándar"), ("Estandar", "Estándar"),
    ("camara", "cámara"), ("camaras", "cámaras"),
    ("resolucion", "resolución"),
    ("linea", "línea"), ("lineas", "líneas"),
    
    # --- Verbos ---
    ("mostro", "mostró"), ("demostro", "demostró"),
    ("detecto", "detectó"), ("codifico", "codificó"),
    ("definio", "definió"), ("presento", "presentó"),
    ("incorporo", "incorporó"), ("aporto", "aportó"),
    ("Colaboro ", "Colaboró "),
    
    # --- Adverbios/conectores ---
    ("tambien", "también"), ("ademas", "además"),
    ("traves", "través"),
    
    # --- Correcciones de las biografías insertadas ---
    ("Ingenieria Industrial", "Ingeniería Industrial"),
    ("areas de interes", "áreas de interés"),
    ("ergonomia industrial", "ergonomía industrial"),
    ("ciberfisicos aplicados", "ciberfísicos aplicados"),
    ("analisis de tiempos", "análisis de tiempos"),
    ("investigacion del", "investigación del"),
    ("sesion de", "sesión de"),
    ("medicion optica", "medición óptica"),
    ("estacion de", "estación de"),
    
    # --- Correcciones del email (revertir daño) ---
    ("ucat\u00f3lica", "ucatolica"),
    ("ucatólica", "ucatolica"),
]

def fix_spelling(text):
    if not text:
        return text
    for old, new in SPELLING:
        text = text.replace(old, new)
    return text

# ============================================================
# Corregir "ptimo/ptima" residual
# ============================================================
def fix_optimo(text):
    if not text:
        return text
    text = re.sub(r'(?<!\w)ptimo\b', 'óptimo', text)
    text = re.sub(r'(?<!\w)ptima\b', 'óptima', text)
    text = re.sub(r'(?<!\w)ptimos\b', 'óptimos', text)
    text = re.sub(r'(?<!\w)ptimas\b', 'óptimas', text)
    return text

# ============================================================
# Pipeline completo
# ============================================================
def apply_all(text):
    text = fix_chars(text)
    text = fix_double(text)
    text = fix_spelling(text)
    text = fix_optimo(text)
    return text

# --- Corregir PARRAFOS ---
changes = 0
for i, para in enumerate(doc.paragraphs):
    for run in para.runs:
        original = run.text
        if original:
            fixed = apply_all(original)
            if fixed != original:
                run.text = fixed
                changes += 1

print(f"Cambios en parrafos: {changes}")

# --- Corregir TABLAS ---
tchn = 0
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    original = run.text
                    if original:
                        fixed = apply_all(original)
                        if fixed != original:
                            run.text = fixed
                            tchn += 1

print(f"Cambios en tablas: {tchn}")

# --- Guardar en archivo temporal ---
doc.save(TEMP_PATH)
print(f"\nDocumento CORREGIDO guardado en: {TEMP_PATH}")
print("Por favor cierra el documento original en Word, luego reemplazalo con el corregido.")
print(f"O simplemente abre: {TEMP_PATH}")
