import docx

file_path = r"c:\Users\dsant\Desktop\U CATOLICA\ING INDUSTRIAL\SEMESTRES INDUSTRIAL\segundo semestre\ingenieria de metodos\Corte 1\ingenieria de metodos\Informe Final Métodos.docx"
doc = docx.Document(file_path)

for idx, p in enumerate(doc.paragraphs):
    has_image = False
    for r in p.runs:
        if r.element.xpath('.//w:drawing') or r.element.xpath('.//w:pict') or r.element.xpath('.//w:object'):
            has_image = True
            break
    if has_image:
        print(f"--- Paragraph {idx} has image ---")
        # print context
        start = max(0, idx - 2)
        end = min(len(doc.paragraphs) - 1, idx + 2)
        for c in range(start, end + 1):
            marker = "-->" if c == idx else "   "
            txt = doc.paragraphs[c].text.strip()
            print(f"{marker} [{c}]: {txt[:100] or '(empty)'}")
