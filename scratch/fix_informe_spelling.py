import docx
import os
import re

file_path = r"c:\Users\dsant\Desktop\U CATOLICA\ING INDUSTRIAL\SEMESTRES INDUSTRIAL\segundo semestre\ingenieria de metodos\Corte 1\ingenieria de metodos\Informe Final Métodos.docx"

if not os.path.exists(file_path):
    print("File not found")
    exit()

doc = docx.Document(file_path)
print(f"Loaded document with {len(doc.paragraphs)} paragraphs and {len(doc.tables)} tables.")

# Corrections dictionary: (unaccented regex pattern, replacement)
# We use word boundaries \b to avoid partial matches
corrections = [
    # General corrections
    (r'\bAnalisis\b', 'Análisis'),
    (r'\banalisis\b', 'análisis'),
    (r'\bBasica\b', 'Básica'),
    (r'\bbasica\b', 'básica'),
    (r'\bBasicos\b', 'Básicos'),
    (r'\bbasicos\b', 'básicos'),
    (r'\bGrafica\b', 'Gráfica'),
    (r'\bgrafica\b', 'gráfica'),
    (r'\bdistribucion\b', 'distribución'),
    (r'\boperacion\b', 'operación'),
    (r'\bmetodologia\b', 'metodología'),
    (r'\bergonomia\b', 'ergonomía'),
    (r'\bergonomico\b', 'ergonómico'),
    (r'\bergonomica\b', 'ergonómica'),
    (r'\bergonomicos\b', 'ergonómicos'),
    (r'\bergonomicas\b', 'ergonómicas'),
    (r'\btecnica\b', 'técnica'),
    (r'\btecnicas\b', 'técnicas'),
    (r'\btecnico\b', 'técnico'),
    (r'\btecnicos\b', 'técnicos'),
    (r'\bmedicion\b', 'medición'),
    (r'\bclasificacion\b', 'clasificación'),
    (r'\bcamara\b', 'cámara'),
    (r'\bcamaras\b', 'cámaras'),
    (r'\bduracion\b', 'duración'),
    (r'\blinea\b', 'línea'),
    (r'\blineas\b', 'líneas'),
    (r'\bposicion\b', 'posición'),
    (r'\bangulo\b', 'ángulo'),
    (r'\bangulos\b', 'ángulos'),
    (r'\bautomatica\b', 'automática'),
    (r'\bautomatico\b', 'automático'),
    (r'\bautomaticamente\b', 'automáticamente'),
    (r'\bdinamico\b', 'dinámico'),
    (r'\bdinamica\b', 'dinámica'),
    (r'\bclasica\b', 'clásica'),
    (r'\bclasico\b', 'clásico'),
    (r'\bcritico\b', 'crítico'),
    (r'\bcritica\b', 'crítica'),
    (r'\binspeccion\b', 'inspección'),
    (r'\bdesviacion\b', 'desviación'),
    (r'\bevaluacion\b', 'evaluación'),
    (r'\boptimo\b', 'óptimo'),
    (r'\boptima\b', 'óptima'),
    (r'\boptimos\b', 'óptimos'),
    (r'\boptimas\b', 'óptimas'),
    (r'\bbasico\b', 'básico'),
    (r'\bbasica\b', 'básica'),
    (r'\bdiseno\b', 'diseño'),
    (r'\bDiseno\b', 'Diseño'),
    (r'\bregresion\b', 'regresión'),
    (r'\bprediccion\b', 'predicción'),
    (r'\bimplementacion\b', 'implementación'),
    (r'\bsolucion\b', 'solución'),
    (r'\bseccion\b', 'sección'),
    (r'\bestimacion\b', 'estimación'),
    (r'\binformacion\b', 'información'),
    (r'\binvestigacion\b', 'investigación'),
    (r'\bcorrelacion\b', 'correlación'),
    (r'\binteraccion\b', 'interacción'),
    (r'\bintegracion\b', 'integración'),
    (r'\bconfiguracion\b', 'configuración'),
    (r'\bcomunicacion\b', 'comunicación'),
    (r'\baplicacion\b', 'aplicación'),
    (r'\bobservacion\b', 'observación'),
    (r'\bmanipulacion\b', 'manipulación'),
    (r'\bactivacion\b', 'activación'),
    (r'\bvisualizacion\b', 'visualización'),
    (r'\bresolucion\b', 'resolución'),
    (r'\btambien\b', 'también'),
    (r'\bademas\b', 'además'),
    (r'\btraves\b', 'través'),
    (r'\bIngeniera\b', 'Ingeniería'),
    (r'\bMtodos\b', 'Métodos'),
    (r'\bergonoma\b', 'ergonomía'),
    (r'\bdeformacion\b', 'deformación'),
    
    # Specific ones found in check_real_typos.py
    (r'\bfisicas\b', 'físicas'),
    (r'\bsegun\b', 'según'),
    (r'\bacrilica\b', 'acrílica'),
    (r'\bmecanica\b', 'mecánica'),
    (r'\binstantanea\b', 'instantánea'),
    (r'\bErgonomico\b', 'Ergonómico'),
    (r'\bflexion\b', 'flexión'),
    (r'\bIluminacion\b', 'Iluminación'),
    (r'\blamparas\b', 'lámparas'),
    (r'\blogro\b', 'logró')
]

def apply_corrections(text):
    if not text:
        return text
    fixed = text
    for pattern, replacement in corrections:
        # Ignore matches inside URLs
        if 'http://' in fixed or 'https://' in fixed:
            # Simple check: if the word is part of the URL domain or path, don't replace it
            # We can split by space and only process tokens that are not URLs
            parts = []
            for token in fixed.split(' '):
                if token.startswith('http://') or token.startswith('https://'):
                    parts.append(token)
                else:
                    parts.append(re.sub(pattern, replacement, token))
            fixed = ' '.join(parts)
        else:
            fixed = re.sub(pattern, replacement, fixed)
    return fixed

# Fix paragraphs
p_changes = 0
for idx, p in enumerate(doc.paragraphs):
    for r in p.runs:
        orig = r.text
        if orig:
            fixed = apply_corrections(orig)
            if fixed != orig:
                r.text = fixed
                p_changes += 1

# Fix tables
t_changes = 0
for t_idx, table in enumerate(doc.tables):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    orig = r.text
                    if orig:
                        fixed = apply_corrections(orig)
                        if fixed != orig:
                            r.text = fixed
                            t_changes += 1

print(f"Applied {p_changes} corrections in paragraphs.")
print(f"Applied {t_changes} corrections in tables.")

doc.save(file_path)
print("Saved corrected document.")
