import docx

file_path = r"c:\Users\dsant\Desktop\U CATOLICA\ING INDUSTRIAL\SEMESTRES INDUSTRIAL\segundo semestre\ingenieria de metodos\Corte 1\ingenieria de metodos\Informe Final Métodos.docx"
doc = docx.Document(file_path)

print("Paragraph 82 text (first 100 chars):", doc.paragraphs[82].text[:100])
if doc.paragraphs[82].runs:
    run = doc.paragraphs[82].runs[0]
    print("Font name:", run.font.name)
    print("Font size:", run.font.size)
    print("Bold:", run.bold)
    print("Italic:", run.italic)
