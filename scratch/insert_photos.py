"""
Script para insertar fotos de los integrantes masculinos del equipo
en el articulo IEEE de CronoGrulla.

Fotos seleccionadas:
- IMG_20260311_164913664.jpg  -> Ambos trabajando juntos (Fig. 2 - puesto experimental)
- IMG_20260311_164720100_HDR.jpg -> Juan Diego operando con la app visible (Fig. cronometro)
- IMG_20260311_164353147.jpg  -> Ambos en el escritorio (foto de autores/equipo)
"""

import os
import sys
import io
from copy import deepcopy
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

# Rutas
DOC_PATH = r"C:\Users\dsant\Desktop\esta\Articulo_CronoGrulla_IEEE.docx"
IMG_DIR = r"C:\Users\dsant\Desktop\U CATOLICA\ING INDUSTRIAL\SEMESTRES INDUSTRIAL\segundo semestre\ingenieria de metodos\Corte 1\ingenieria de metodos\Imagenes"

# Fotos seleccionadas de los hombres
PHOTO_BOTH = os.path.join(IMG_DIR, "IMG_20260311_164913664.jpg")       # Ambos juntos close-up
PHOTO_JUAN_DIEGO = os.path.join(IMG_DIR, "IMG_20260311_164720100_HDR.jpg")  # Juan Diego solo con app
PHOTO_TEAM = os.path.join(IMG_DIR, "IMG_20260311_164353147.jpg")       # Ambos en escritorio

# Verificar que existen
for f in [PHOTO_BOTH, PHOTO_JUAN_DIEGO, PHOTO_TEAM]:
    if not os.path.exists(f):
        print(f"ERROR: No se encuentra {f}")
        sys.exit(1)

print("Todas las fotos encontradas OK.")

# Abrir documento
doc = Document(DOC_PATH)
print(f"Documento abierto: {len(doc.paragraphs)} parrafos")

def find_paragraph_index(doc, search_text):
    """Buscar indice de un parrafo que contenga cierto texto."""
    for i, para in enumerate(doc.paragraphs):
        if search_text.lower() in para.text.lower():
            return i
    return -1

def insert_image_after_paragraph(doc, para_index, image_path, caption_text, width_inches=3.2):
    """Insertar una imagen con caption despues de un parrafo especifico."""
    # Obtener el elemento XML del parrafo de referencia
    ref_para = doc.paragraphs[para_index]
    ref_element = ref_para._element
    
    # Crear nuevo parrafo para la imagen
    new_img_para = doc.add_paragraph()
    new_img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = new_img_para.runs[0] if new_img_para.runs else new_img_para.add_run()
    run.add_picture(image_path, width=Inches(width_inches))
    
    # Crear nuevo parrafo para el caption
    new_cap_para = doc.add_paragraph()
    new_cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = new_cap_para.add_run(caption_text)
    cap_run.font.size = Pt(8)
    cap_run.font.italic = True
    
    # Mover los elementos despues del parrafo de referencia
    # (se agregaron al final, hay que moverlos)
    img_element = new_img_para._element
    cap_element = new_cap_para._element
    
    # Remover de su posicion actual (al final del body)
    doc.element.body.remove(img_element)
    doc.element.body.remove(cap_element)
    
    # Insertar despues del parrafo de referencia
    ref_element.addnext(cap_element)
    ref_element.addnext(img_element)
    
    print(f"  Imagen insertada despues del parrafo [{para_index}]")
    print(f"  Caption: {caption_text}")

# === INSERCION 1: Foto del puesto experimental (Fig. 2) ===
# Buscar "Fig. 2. Puesto experimental"
fig2_idx = find_paragraph_index(doc, "Fig. 2.")
if fig2_idx >= 0:
    print(f"\n[1] Encontrado 'Fig. 2' en parrafo [{fig2_idx}]")
    # Insertar la foto de ambos en el puesto de trabajo ANTES del caption de Fig. 2
    # Reemplazar el caption de Fig. 2 con la imagen real
    insert_image_after_paragraph(
        doc, fig2_idx - 1,  # Insertar antes del caption
        PHOTO_BOTH,
        "Fig. 2. Puesto experimental de manufactura de origami con medicion optica directa en la estacion de trabajo.",
        width_inches=3.3
    )
    # Eliminar el caption viejo de Fig. 2
    old_fig2 = doc.paragraphs[fig2_idx + 2]  # +2 porque insertamos 2 parrafos antes
    if "Fig. 2" in old_fig2.text:
        old_fig2._element.getparent().remove(old_fig2._element)
        print("  Caption viejo de Fig. 2 eliminado")
else:
    print("\n[1] No se encontro 'Fig. 2', insertando al final de seccion III")
    sec3_idx = find_paragraph_index(doc, "IV. RESULTADOS")
    if sec3_idx >= 0:
        insert_image_after_paragraph(
            doc, sec3_idx - 1,
            PHOTO_BOTH,
            "Fig. 2. Puesto experimental de manufactura de origami con medicion optica directa en la estacion de trabajo.",
            width_inches=3.3
        )

# === INSERCION 2: Foto de Juan Diego operando el cronometro ===
# Buscar la seccion de resultados para insertar la foto operativa
fig3_idx = find_paragraph_index(doc, "Fig. 3.")
if fig3_idx >= 0:
    print(f"\n[2] Encontrado 'Fig. 3' en parrafo [{fig3_idx}]")
    insert_image_after_paragraph(
        doc, fig3_idx - 1,
        PHOTO_JUAN_DIEGO,
        "Fig. 3. Operario ejecutando el plegado de origami con el sistema CronoGrulla registrando tiempos en tiempo real.",
        width_inches=3.3
    )
    # Eliminar caption viejo
    old_fig3 = doc.paragraphs[fig3_idx + 2]
    if "Fig. 3" in old_fig3.text:
        old_fig3._element.getparent().remove(old_fig3._element)
        print("  Caption viejo de Fig. 3 eliminado")
else:
    print("\n[2] No se encontro 'Fig. 3', insertando despues de resultados")
    res_idx = find_paragraph_index(doc, "IV. RESULTADOS")
    if res_idx >= 0:
        insert_image_after_paragraph(
            doc, res_idx + 2,
            PHOTO_JUAN_DIEGO,
            "Fig. 3. Operario ejecutando el plegado de origami con el sistema CronoGrulla registrando tiempos en tiempo real.",
            width_inches=3.3
        )

# === INSERCION 3: Foto del equipo al final (seccion de autores IEEE) ===
# En formato IEEE, las fotos de autores van al final, despues de las referencias
ref_idx = find_paragraph_index(doc, "REFERENCIAS")
last_ref_idx = -1
for i, para in enumerate(doc.paragraphs):
    if para.text.strip().startswith("[9]"):
        last_ref_idx = i
        break

if last_ref_idx >= 0:
    print(f"\n[3] Ultima referencia encontrada en parrafo [{last_ref_idx}]")
    
    # Agregar seccion de biografias de autores (estilo IEEE)
    # Primero agregar un separador
    bio_sep = doc.add_paragraph()
    bio_sep.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Foto del equipo trabajando
    bio_img = doc.add_paragraph()
    bio_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bio_run = bio_img.add_run()
    bio_run.add_picture(PHOTO_TEAM, width=Inches(3.3))
    
    # Caption de la foto del equipo
    bio_cap = doc.add_paragraph()
    bio_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = bio_cap.add_run("Fig. 6. Equipo de investigacion del proyecto CronoGrulla durante la sesion de toma de datos experimentales.")
    cap_run.font.size = Pt(8)
    cap_run.font.italic = True
    
    # Bio de David Santiago
    bio_david = doc.add_paragraph()
    bio_david.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    dr = bio_david.add_run("David Santiago Castelblanco Artunduaga ")
    dr.font.bold = True
    dr.font.size = Pt(9)
    dr2 = bio_david.add_run("es estudiante de Ingenieria Industrial en la Universidad Catolica de Colombia. Sus areas de interes incluyen la ingenieria de metodos, la ergonomia industrial y los sistemas ciberfisicos aplicados a la manufactura.")
    dr2.font.size = Pt(9)
    
    # Bio de Juan Diego
    bio_juan = doc.add_paragraph()
    bio_juan.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    jr = bio_juan.add_run("Juan Diego Escobar Duarte ")
    jr.font.bold = True
    jr.font.size = Pt(9)
    jr2 = bio_juan.add_run("es estudiante de Ingenieria Industrial en la Universidad Catolica de Colombia. Colaboro en la toma de datos experimentales y el analisis de tiempos y micromovimientos del proceso de manufactura de origami.")
    jr2.font.size = Pt(9)
    
    print("  Seccion de autores con foto insertada al final del documento")
else:
    print("\n[3] No se encontro la ultima referencia")

# Guardar
backup_path = DOC_PATH.replace(".docx", "_backup_sin_fotos.docx")
import shutil
if not os.path.exists(backup_path):
    shutil.copy2(DOC_PATH, backup_path)
    print(f"\nBackup creado: {backup_path}")

doc.save(DOC_PATH)
print(f"\nDocumento guardado exitosamente en: {DOC_PATH}")
print("Fotos de los hombres del equipo insertadas correctamente.")
