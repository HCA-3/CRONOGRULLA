import docx

file_path = r"c:\Users\dsant\Desktop\U CATOLICA\ING INDUSTRIAL\SEMESTRES INDUSTRIAL\segundo semestre\ingenieria de metodos\Corte 1\ingenieria de metodos\Codigo.docx"
doc = docx.Document(file_path)

p = doc.paragraphs[0]
print("P[0] text:", p.text)
for c in p.text:
    print(f"'{c}': {ord(c)}")
    
p19 = doc.paragraphs[19]
print("\nP[19] text:", p19.text)
for c in p19.text:
    print(f"'{c}': {ord(c)}")
