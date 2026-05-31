import docx

file_path = r"c:\Users\dsant\Desktop\U CATOLICA\ING INDUSTRIAL\SEMESTRES INDUSTRIAL\segundo semestre\ingenieria de metodos\Corte 1\ingenieria de metodos\Informe Final Métodos.docx"
doc = docx.Document(file_path)

print("Paragraph 81 (Conclusiones) text:", doc.paragraphs[81].text)
print("Paragraph 81 style:", doc.paragraphs[81].style.name)
if doc.paragraphs[81].runs:
    run = doc.paragraphs[81].runs[0]
    print("Font name:", run.font.name)
    print("Font size:", run.font.size)
    print("Bold:", run.bold)
    print("Color:", run.font.color.rgb if run.font.color else "No color")

print("\nParagraph 87 (Anexos) text:", doc.paragraphs[87].text)
print("Paragraph 87 style:", doc.paragraphs[87].style.name)
if doc.paragraphs[87].runs:
    run = doc.paragraphs[87].runs[0]
    print("Font name:", run.font.name)
    print("Font size:", run.font.size)
    print("Bold:", run.bold)
