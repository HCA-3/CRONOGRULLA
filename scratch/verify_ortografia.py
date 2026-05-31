"""Verificar el documento corregido y analizar las tablas."""
import docx
import sys
import io

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

p = r"C:\Users\dsant\Desktop\esta\Articulo_CronoGrulla_IEEE_CORREGIDO.docx"
doc = docx.Document(p)

print("=== VERIFICAR PARRAFOS CLAVE ===\n")
keywords = ['est', 'ptimo', 'ptima', 'cion', 'ologia', 'ergon', 'metod', 'angul', 
            'indice', 'articulo', 'investigacion', 'sesion', 'ucatolica']
for i, para in enumerate(doc.paragraphs):
    text = para.text
    if not text.strip():
        continue
    # Buscar palabras sin tildes que deberian tenerlas
    lower = text.lower()
    issues = []
    
    # Verificar patrones problematicos
    if 'estáá' in text or 'est??' in text:
        issues.append("DOBLE CORRUPCION")
    if ' ptimo' in lower or ' ptima' in lower:
        issues.append("OPTIMO SIN TILDE")
    if 'metodologica' in lower and 'metodológica' not in lower:
        issues.append("METODOLOGICA")
    
    if issues:
        safe = text[:100].encode('ascii', 'replace').decode('ascii')
        print(f"  [{i}] PROBLEMAS: {issues}")
        print(f"       {safe}")
        print()

print("\n=== TABLAS - TEXTO RUN POR RUN ===\n")
for t, table in enumerate(doc.tables):
    print(f"--- TABLA {t} ---")
    for r, row in enumerate(table.rows):
        for c, cell in enumerate(row.cells):
            for p, para in enumerate(cell.paragraphs):
                runs_text = []
                for run in para.runs:
                    if run.text:
                        # Mostrar caracteres problematicos
                        for ch in run.text:
                            code = ord(ch)
                            if code > 127:
                                runs_text.append(f"[chr({code})={ch}]")
                            else:
                                runs_text.append(ch)
                full = "".join(runs_text)
                if full.strip():
                    print(f"  [{r},{c}] {full}")
    print()

print("\n=== MUESTRA DE PARRAFOS CORREGIDOS ===\n")
sample = [0, 1, 3, 10, 12, 33, 51, 76, 77, 78]
for i in sample:
    if i < len(doc.paragraphs):
        text = doc.paragraphs[i].text
        if text.strip():
            safe = text[:150].encode('ascii', 'replace').decode('ascii')
            print(f"  [{i}] {safe}")
