import docx
import os
import re

file_path = r"c:\Users\dsant\Desktop\U CATOLICA\ING INDUSTRIAL\SEMESTRES INDUSTRIAL\segundo semestre\ingenieria de metodos\Corte 1\ingenieria de metodos\Informe Final Métodos.docx"

if not os.path.exists(file_path):
    print("File not found")
    exit()

doc = docx.Document(file_path)

# Only search for actual spelling errors (unaccented words)
real_typos = [
    r'\bmetodologia\b', r'\bmetodologias\b', r'\bproduccion\b', r'\bcalificacion\b',
    r'\bestandar\b', r'\bdefinicion\b', r'\btecnica\b', r'\btecnicas\b', 
    r'\bmedicion\b', r'\bmediciones\b', r'\bclasificacion\b', r'\bclasificaciones\b',
    r'\banalisis\b', r'\bgrafico\b', r'\bgraficos\b', r'\bcamara\b', r'\bcamaras\b', 
    r'\bduracion\b', r'\blinea\b', r'\blineas\b', r'\bergonomia\b', r'\bergonomico\b',
    r'\bergonomica\b', r'\bergonomicos\b', r'\bergonomicas\b', r'\bptimo\b',
    r'\bptima\b', r'\bptimos\b', r'\bptimas\b', r'\bdesviacion\b', r'\bevaluacion\b',
    r'\boptimo\b', r'\boptima\b', r'\boptimos\b', r'\boptimas\b', r'\bbasico\b',
    r'\bbasica\b', r'\bbasicos\b', r'\bbasicas\b', r'\bregresion\b', r'\bprediccion\b',
    r'\bimplementacion\b', r'\bsolucion\b', r'\bseccion\b', r'\bestimacion\b',
    r'\binformacion\b', r'\binvestigacion\b', r'\bcorrelacion\b', r'\binteraccion\b',
    r'\bintegracion\b', r'\bconfiguracion\b', r'\bcomunicacion\b', r'\baplicacion\b',
    r'\bobservacion\b', r'\bmanipulacion\b', r'\bactivacion\b', r'\bvisualizacion\b',
    r'\bresolucion\b', r'\btambien\b', r'\bademas\b', r'\btraves\b', r'\bIngeniera\b',
    r'\bMtodos\b', r'\bergonoma\b'
]

print("\n--- Checking paragraphs for actual spelling errors ---")
count = 0
for idx, p in enumerate(doc.paragraphs):
    text = p.text
    for pattern in real_typos:
        if re.search(pattern, text, re.IGNORECASE):
            print(f"P[{idx}] Typo '{pattern}': {text}")
            count += 1
            break
            
print(f"Found {count} real typos in paragraphs.")

print("\n--- Checking tables for actual spelling errors ---")
tbl_count = 0
for t_idx, table in enumerate(doc.tables):
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            text = cell.text
            for pattern in real_typos:
                if re.search(pattern, text, re.IGNORECASE):
                    print(f"Table[{t_idx}] Cell[{r_idx},{c_idx}] Typo '{pattern}': {text}")
                    tbl_count += 1
                    break
print(f"Found {tbl_count} real typos in tables.")
