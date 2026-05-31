import docx
import os

file_path = r"c:\Users\dsant\Desktop\U CATOLICA\ING INDUSTRIAL\SEMESTRES INDUSTRIAL\segundo semestre\ingenieria de metodos\Corte 1\ingenieria de metodos\Informe Final Métodos.docx"

try:
    doc = docx.Document(file_path)
    print("Successfully opened!")
    print("Number of paragraphs:", len(doc.paragraphs))
    print("Number of tables:", len(doc.tables))
    print("\n--- Last 30 paragraphs ---")
    start = max(0, len(doc.paragraphs) - 30)
    for i in range(start, len(doc.paragraphs)):
        print(f"{i}: {doc.paragraphs[i].text}")
except Exception as e:
    print("Error:", e)
