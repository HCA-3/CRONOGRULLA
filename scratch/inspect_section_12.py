import docx

file_path = r"c:\Users\dsant\Desktop\U CATOLICA\ING INDUSTRIAL\SEMESTRES INDUSTRIAL\segundo semestre\ingenieria de metodos\Corte 1\ingenieria de metodos\Informe Final Métodos.docx"
doc = docx.Document(file_path)

found = False
for idx, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    if "12." in txt or "therblig" in txt.lower() or "micro-movimientos" in txt.lower() or "micromovimientos" in txt.lower():
        print(f"Index {idx}: {txt}")
        found = True

if not found:
    print("No direct match found in paragraph text. Searching runs or substring...")
    for idx, p in enumerate(doc.paragraphs):
        if any(term in p.text.lower() for term in ["micro", "therblig"]):
            print(f"Partial Index {idx}: {p.text[:120]}")
