"""
Script para corregir las tablas del documento IEEE.
Problema: 'Óóóóóptimo' debe ser 'Óptimo'
"""
import docx
import sys
import io
import re

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

DOC_PATH = r"C:\Users\dsant\Desktop\esta\Articulo_CronoGrulla_IEEE_CORREGIDO.docx"
doc = docx.Document(DOC_PATH)

print("=== CORRIGIENDO TABLAS ===\n")

changes = 0
for t, table in enumerate(doc.tables):
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    text = run.text
                    if not text:
                        continue
                    original = text
                    
                    # Patron: Ó seguido de múltiples ó + ptimo -> Óptimo
                    text = re.sub(r'Ó[óÓ]+ptimo', 'Óptimo', text)
                    text = re.sub(r'Ó[óÓ]+ptima', 'Óptima', text)
                    
                    # Patron: ó seguido de múltiples ó + ptimo -> óptimo
                    text = re.sub(r'ó[óÓ]+ptimo', 'óptimo', text)
                    text = re.sub(r'ó[óÓ]+ptima', 'óptima', text)
                    
                    # Corregir "metodológica" que quedó como "métodológica"
                    text = text.replace('métodológica', 'metodológica')
                    text = text.replace('Métodológica', 'Metodológica')
                    
                    if text != original:
                        run.text = text
                        changes += 1
                        o = original.encode('ascii','replace').decode('ascii')
                        n = text.encode('ascii','replace').decode('ascii')
                        print(f"  Tabla {t}: '{o}' -> '{n}'")

print(f"\nCambios en tablas: {changes}")

# Tambien corregir en parrafos las mismas cosas
pchanges = 0
for i, para in enumerate(doc.paragraphs):
    for run in para.runs:
        text = run.text
        if not text:
            continue
        original = text
        
        # Corregir "metodológica" duplicada
        text = text.replace('métodológica', 'metodológica')
        text = text.replace('Métodológica', 'Metodológica')
        text = text.replace('métodológicas', 'metodológicas')
        
        # Corregir "métodos" que pudo quedar como "métodos" (ya correcto) 
        # pero "metodológica" pudo quedar con doble tilde
        
        # Corregir "composiciónal" -> "composicional"
        text = text.replace('composiciónal', 'composicional')
        
        # Corregir "acciónes" -> "acciones"  
        text = text.replace('acciónes', 'acciones')
        
        # Corregir "operaciónes" -> "operaciones"
        text = text.replace('operaciónes', 'operaciones')
        
        # Corregir "investigaciónes" -> "investigaciones"
        text = text.replace('investigaciónes', 'investigaciones')
        
        # Patron general: -iónes -> -iones (plural no lleva tilde en -iones)
        text = re.sub(r'iónes\b', 'iones', text)
        
        if text != original:
            run.text = text
            pchanges += 1
            o = original[:80].encode('ascii','replace').decode('ascii')
            n = text[:80].encode('ascii','replace').decode('ascii')
            print(f"  P[{i}]: '{o}' -> '{n}'")

print(f"Cambios en parrafos: {pchanges}")

doc.save(DOC_PATH)
print(f"\nGuardado: {DOC_PATH}")
