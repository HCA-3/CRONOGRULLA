import docx
import sys
import io

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

p = r"C:\Users\dsant\Desktop\esta\Articulo_CronoGrulla_IEEE.docx"
doc = docx.Document(p)

from docx.oxml.ns import qn

print("=== VERIFICACION DEL DOCUMENTO ===")
print(f"Total parrafos: {len(doc.paragraphs)}")

# Contar imagenes
img_count = 0
for rel in doc.part.rels.values():
    if "image" in rel.reltype:
        img_count += 1
        print(f"  Imagen encontrada: {rel.target_ref}")
print(f"\nTotal imagenes en el documento: {img_count}")

# Verificar figuras
print("\n=== FIGURAS ===")
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text.startswith("Fig."):
        safe = text[:100].encode('ascii', 'replace').decode('ascii')
        print(f"  [{i}] {safe}")

# Verificar seccion de autores al final
print("\n=== ULTIMOS 10 PARRAFOS ===")
total = len(doc.paragraphs)
for i in range(max(0, total - 10), total):
    text = doc.paragraphs[i].text.strip()
    if text:
        safe = text[:100].encode('ascii', 'replace').decode('ascii')
        print(f"  [{i}] {safe}")

# Verificar imagenes inline
inline_count = 0
for i, para in enumerate(doc.paragraphs):
    drawings = para._element.findall(f'.//{qn("w:drawing")}')
    if drawings:
        text = para.text[:40].encode('ascii','replace').decode('ascii') if para.text else "(vacio)"
        print(f"\n  Imagen inline en parrafo [{i}]: {text}")
        inline_count += 1
print(f"\nTotal imagenes inline: {inline_count}")

print("\nVerificacion completada OK!")
