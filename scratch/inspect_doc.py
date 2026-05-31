import docx
import sys
import io

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

p = r"C:\Users\dsant\Desktop\esta\Articulo_CronoGrulla_IEEE.docx"
doc = docx.Document(p)

print("=== PARRAFOS DEL DOCUMENTO ===")
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text:
        style = para.style.name if para.style else "None"
        safe = text[:120].encode('ascii', 'replace').decode('ascii')
        print(f"[{i}] Style: {style} | {safe}")

print(f"\nTotal paragraphs: {len(doc.paragraphs)}")

print("\n=== BUSCAR SECCION AUTORES ===")
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip().lower()
    if 'david' in text or 'juan' in text or 'autor' in text or 'castelblanco' in text or 'escobar' in text:
        safe = para.text.strip()[:150].encode('ascii', 'replace').decode('ascii')
        print(f"  [{i}] {safe}")

print("\n=== IMAGENES EXISTENTES ===")
img_count = 0
for rel in doc.part.rels.values():
    if "image" in rel.reltype:
        img_count += 1
        print(f"  Imagen: {rel.target_ref}")
print(f"Total imagenes: {img_count}")

# Check for inline shapes (embedded images)
from docx.oxml.ns import qn
inline_count = 0
for i, para in enumerate(doc.paragraphs):
    drawings = para._element.findall(f'.//{qn("w:drawing")}')
    if drawings:
        print(f"  Inline image in paragraph [{i}]: {para.text[:60].encode('ascii','replace').decode('ascii')}")
        inline_count += 1
print(f"Total inline images: {inline_count}")
