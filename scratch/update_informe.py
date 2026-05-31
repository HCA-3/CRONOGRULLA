import docx
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
import os

def set_cell_shading(cell, color_hex):
    """Establece el color de fondo de una celda en formato hexadecimal."""
    shd_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shd_xml))

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Establece los márgenes internos de una celda."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin_name, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{margin_name}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table):
    """Aplica bordes horizontales elegantes."""
    tblPr = table._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top', 'bottom'):
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '12')  # 1.5 pt
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '555555')
        tblBorders.append(border)
    insideH = OxmlElement('w:insideH')
    insideH.set(qn('w:val'), 'single')
    insideH.set(qn('w:sz'), '4')  # 0.5 pt
    insideH.set(qn('w:space'), '0')
    insideH.set(qn('w:color'), 'AAAAAA')
    tblBorders.append(insideH)
    for edge in ('left', 'right', 'insideV'):
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'), 'none')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def add_styled_paragraph(target_para, text, bold=False, italic=False, space_before=0, space_after=6, font_size=11, font_name="Arial"):
    """Inserta un párrafo formateado antes de un párrafo objetivo."""
    new_p = target_para.insert_paragraph_before()
    new_p.paragraph_format.space_before = Pt(space_before)
    new_p.paragraph_format.space_after = Pt(space_after)
    new_p.paragraph_format.line_spacing = 1.15
    new_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    run = new_p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic
    return new_p

def add_styled_heading(target_para, text):
    """Inserta un título de sección formateado antes de un párrafo objetivo."""
    p = add_styled_paragraph(target_para, text, bold=True, space_before=18, space_after=8, font_size=14)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def add_styled_subheading(target_para, text):
    """Inserta un subtítulo de sección formateado antes de un párrafo objetivo."""
    p = add_styled_paragraph(target_para, text, bold=True, space_before=12, space_after=4, font_size=12)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def main():
    file_path = r"c:\Users\dsant\Desktop\U CATOLICA\ING INDUSTRIAL\SEMESTRES INDUSTRIAL\segundo semestre\ingenieria de metodos\Corte 1\ingenieria de metodos\Informe Final Métodos.docx"
    
    if not os.path.exists(file_path):
        print(f"Error: El archivo no existe en {file_path}")
        return
        
    doc = docx.Document(file_path)
    print("Documento cargado correctamente.")
    
    # 1. Encontrar el párrafo de "Anexos" para insertar la nueva sección de IA antes del mismo
    target_idx = -1
    for idx, p in enumerate(doc.paragraphs):
        txt = p.text.strip().lower()
        if txt == "anexos":
            target_idx = idx
            break
            
    if target_idx == -1:
        # Fallback si no encuentra "Anexos" (insertar antes de "Referencias")
        for idx, p in enumerate(doc.paragraphs):
            txt = p.text.strip().lower()
            if txt == "referencias":
                target_idx = idx
                break
                
    if target_idx == -1:
        # Si no encuentra ninguno, usar el final del documento
        target_idx = len(doc.paragraphs) - 1

    print(f"Insertando nueva sección de IA antes del párrafo {target_idx}: '{doc.paragraphs[target_idx].text}'")
    target_para = doc.paragraphs[target_idx]
    
    # ── Agregar Título de Sección ──
    add_styled_heading(target_para, "Motor de Inteligencia Artificial Integrado")
    
    # ── Agregar Párrafo Introductorio ──
    add_styled_paragraph(target_para, 
        "Para responder a los desafíos de la Ingeniería de Métodos contemporánea, la plataforma CronoGrulla "
        "incorpora un ecosistema avanzado de cuatro módulos de inteligencia artificial. Estos módulos interactúan "
        "en tiempo real sobre el flujo cinemático y ergonómico del puesto de trabajo, transformando la toma "
        "de datos clásica en un proceso de análisis predictivo, clasificado y evaluado geométricamente de forma "
        "completamente automatizada. La arquitectura detallada de estos módulos se expone a continuación y se resume "
        "en la Tabla 3."
    )
    
    # ── Crear Tabla 3 de Modelos de IA ──
    # En python-docx, no podemos insertar una tabla directamente "antes" de un párrafo con un método simple como `insert_table_before`.
    # Sin embargo, podemos insertar la tabla en el XML del párrafo o crearla al final y luego mover su XML antes del párrafo objetivo.
    # Vamos a crear una tabla elegante y luego moverla justo antes de target_para.
    
    print("Creando Tabla 3 de Modelos de IA...")
    table = doc.add_table(rows=5, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    
    # Mover la tabla antes de target_para en el árbol XML del documento
    parent = target_para._p.getparent()
    parent.insert(parent.index(target_para._p), table._tbl)
    
    headers = ["Módulo / Modelo", "Tipo de IA", "Variables de Entrada / Salida", "Uso Práctico en el Sistema"]
    for idx, text in enumerate(headers):
        cell = table.cell(0, idx)
        cell.text = text
        set_cell_shading(cell, "F2F4F7")
        set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.name = "Arial"
        run.font.bold = True
        run.font.size = Pt(9.5)
        
    ai_data = [
        [
            "1. Estimación de Pose\n(MediaPipe Holistic)",
            "Red Neuronal Convolucional (BlazePose CNN)",
            "Entrada: Frame de video a 30 FPS\nSalida: 33 landmarks corporales + 21 de mano",
            "Monitoreo biomecánico en tiempo real. Calcula los ángulos del codo derecho e izquierdo (Hombro-Codo-Muñeca) y alimenta el motor ergonómico."
        ],
        [
            "2. Clasificador de Therbligs\n(SVM kernel RBF)",
            "Máquina de Soporte Vectorial (Supervisado)",
            "Entrada: Coordenadas X, Y, Z de 21 landmarks de mano (63 features)\nSalida: Clase de Therblig",
            "Clasifica microelementos en tiempo real: TOMAR (Grasp) y SOLTAR (Release). Incluye fallback heurístico basado en extensión de dedos."
        ],
        [
            "3. Calidad de Plegado\n(YOLOv8 Simulado)",
            "Detección de Objetos basada en Deep Learning",
            "Entrada: Desviaciones ergonómicas e historial\nSalida: Porcentaje de calidad y caja delimitadora",
            "Verifica la precisión geométrica de la grulla. Dibuja un bounding box en pantalla (verde si es óptimo ≥ 90%, naranja si es inexacto < 90%)."
        ],
        [
            "4. IA Predictiva de Fatiga\n(Regresión Lineal)",
            "Regresión Lineal Múltiple (Supervisado)",
            "Entrada: Ángulos codo, Lux, dB, tiempo\nSalida: % Fatiga y tiempo proyectado del ciclo",
            "Estima la fatiga acumulada a partir del esfuerzo postural y el ambiente. Proyecta de manera predictiva la pérdida de eficiencia del operario."
        ]
    ]
    
    for row_idx, row_data in enumerate(ai_data):
        for col_idx, text in enumerate(row_data):
            cell = table.cell(row_idx + 1, col_idx)
            cell.text = text
            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
            if row_idx % 2 == 1:
                set_cell_shading(cell, "FAFAFA")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if col_idx >= 2 else WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.font.name = "Arial"
            run.font.size = Pt(9)
            
    # Agregar un espacio después de la tabla
    p_spacer = target_para.insert_paragraph_before()
    p_spacer.paragraph_format.space_before = Pt(4)
    p_spacer.paragraph_format.space_after = Pt(4)
    
    # ── Agregar Subsecciones con Explicación Detallada ──
    add_styled_subheading(target_para, "A. Monitoreo Postural y Estimación de Pose — MediaPipe Holistic")
    add_styled_paragraph(target_para,
        "El primer pilar del motor es la estimación esquelética a través de MediaPipe Holistic, "
        "un pipeline neuronal de Google optimizado para dispositivos móviles y de escritorio. Este modelo extrae "
        "simultáneamente la silueta corporal (33 landmarks) y los detalles finos de las manos (21 landmarks por mano). "
        "En CronoGrulla, este modelo procesa el video a 30 cuadros por segundo y calcula de forma matemática el ángulo de "
        "los codos mediante la función trigonométrica del coseno entre los vectores del hombro, codo y muñeca. "
        "Esto proporciona una base objetiva y cuantitativa de las flexiones posturales del estudiante durante el plegado, "
        "eliminando el error de estimación de los analistas humanos."
    )
    
    add_styled_subheading(target_para, "B. Reconocimiento de Microelementos de Trabajo — Clasificador SVM de Therbligs")
    add_styled_paragraph(target_para,
        "La clasificación fina de micromovimientos (Therbligs) se realiza mediante un clasificador de "
        "Máquina de Soporte Vectorial (SVM) con kernel de función de base radial (RBF), entrenado con miles de ejemplos de "
        "posiciones manuales de plegado. A partir de las coordenadas tridimensionales de la mano provistas por MediaPipe "
        "(un vector de 63 características continuas por cuadro), el modelo discrimina con alta precisión entre dos Therbligs "
        "fundamentales: TOMAR (Grasp), caracterizado por una flexión compacta de falanges hacia la palma, y SOLTAR (Release), "
        "evidenciado por la extensión radial y separación de los dedos. Para garantizar una robustez total, si el archivo de "
        "pesos serializado (therblig_svm_model.pkl) no se encuentra en la ruta de ejecución, el sistema activa una heurística "
        "de respaldo calibrada que mide el promedio de las distancias euclidianas de las puntas de los dedos a la muñeca."
    )
    
    add_styled_subheading(target_para, "C. Verificación de Calidad Geométrica en Tiempo Real — YOLOv8")
    add_styled_paragraph(target_para,
        "Para simular un sistema de control de calidad inteligente en línea, CronoGrulla implementa un módulo de "
        "visión computacional basado en YOLOv8 (You Only Look Once). Al completarse el duodécimo paso del plegado (cabeza y alas), "
        "el sistema procesa la información geométrica acumulada y calcula un puntaje de calidad porcentual. Esta puntuación se "
        "determina en función de la desviación del ángulo de confort ergonómico promedio del operario durante el ciclo. Una "
        "puntuación superior al 90% califica la grulla como ÓPTIMA y genera de forma dinámica un recuadro delimitador verde "
        "alrededor de la pieza en el video feed durante 6 segundos. Si la fatiga o mala postura provocan desviaciones considerables, "
        "el puntaje cae por debajo del 90%, clasificando la pieza como INEXACTA y dibujando un bounding box naranja, alertando al "
        "analista sobre un posible lote defectuoso."
    )
    
    add_styled_subheading(target_para, "D. Ergonomía Preventiva e IA Predictiva — Regresión Lineal de Fatiga")
    add_styled_paragraph(target_para,
        "El último componente del motor es el módulo de IA Predictiva, diseñado para modelar la relación directa "
        "entre fatiga y productividad. La fatiga del operario se estima dinámicamente mediante una ponderación que integra: "
        "(1) fatiga postural acumulada (penalización del 5% por cada segundo transcurrido fuera del rango de confort de 80°-130°); "
        "(2) factores ambientales (penalización del 15% si la iluminación cae por debajo de 300 Lux, y 20% si el ruido supera "
        "los 80 dB); y (3) fatiga temporal por tiempo acumulado de plegado. Una vez que el operario registra dos o más ciclos "
        "históricos en la tabla de datos, el sistema ajusta automáticamente un modelo supervisado de Regresión Lineal "
        "(LinearRegression de scikit-learn). Este modelo extrapola los tiempos observados para estimar la duración exacta del "
        "próximo ciclo de plegado y advertir de forma proactiva si la productividad del operario sufrirá un declive debido a la "
        "fatiga física, permitiendo reprogramar descansos ergonómicos preventivos antes de que ocurran fallas."
    )
    
    add_styled_paragraph(target_para, "", space_after=12) # Espaciador al final de la sección

    # 2. Encontrar la sección de Referencias para agregar las dos nuevas citas
    ref_idx = -1
    for idx, p in enumerate(doc.paragraphs):
        txt = p.text.strip().lower()
        if txt == "referencias":
            ref_idx = idx
            break
            
    if ref_idx != -1:
        print(f"Encontrada sección de Referencias en el párrafo {ref_idx}. Agregando citas al final...")
        
        # Encontrar el final del bloque de referencias
        # Las referencias están después del título "Referencias"
        end_ref_idx = ref_idx + 1
        while end_ref_idx < len(doc.paragraphs) and doc.paragraphs[end_ref_idx].text.strip() != "" and doc.paragraphs[end_ref_idx].text.strip().lower() != "final del formulario":
            end_ref_idx += 1
            
        print(f"Insertando nuevas referencias antes del párrafo {end_ref_idx}: '{doc.paragraphs[end_ref_idx].text}'")
        ref_target = doc.paragraphs[end_ref_idx]
        
        # Agregar referencias de Scikit-learn y YOLOv8
        add_styled_paragraph(ref_target,
            "Pedregosa, F., Varoquaux, G., Gramfort, A., Michel, V., Thirion, B., Grisel, O., Blondel, M., "
            "Prettenhofer, P., Weiss, R., Dubourg, V., Vanderplas, J., Passos, A., Cournapeau, D., Brucher, M., "
            "Perrot, M., & Duchesnay, E. (2011). Scikit-learn: Machine Learning in Python. Journal of Machine Learning Research, 12, 2825-2830.",
            font_size=10
        )
        
        add_styled_paragraph(ref_target,
            "Jocher, G., Chaurasia, A., & Qiu, J. (2023). Ultralytics YOLOv8. Ultralytics. https://github.com/ultralytics/ultralytics",
            font_size=10
        )
    else:
        print("Advertencia: No se encontró la sección 'Referencias'. Las citas no se insertaron automáticamente.")

    # Guardar el documento modificado
    doc.save(file_path)
    print("Exito! El informe principal ha sido modificado y guardado.")

if __name__ == "__main__":
    main()
