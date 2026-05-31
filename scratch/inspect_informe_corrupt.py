import docx
import os

file_path = r"c:\Users\dsant\Desktop\U CATOLICA\ING INDUSTRIAL\SEMESTRES INDUSTRIAL\segundo semestre\ingenieria de metodos\Corte 1\ingenieria de metodos\Informe Final Métodos.docx"

if not os.path.exists(file_path):
    print("File not found")
    exit()

doc = docx.Document(file_path)
print(f"Total Paragraphs: {len(doc.paragraphs)}")

count = 0
for idx, p in enumerate(doc.paragraphs):
    text = p.text
    if '' in text:
        print(f"P[{idx}]: {text}")
        count += 1

print(f"\nFound {count} paragraphs with ''.")
