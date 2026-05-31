import os
import docx

file_path = r"C:\Users\dsant\Desktop\esta\PLANTILLA-IEEE_2026.doc"

try:
    # Try reading as docx
    doc = docx.Document(file_path)
    print("Document successfully opened with python-docx!")
    print("Number of paragraphs:", len(doc.paragraphs))
    print("\n--- First 20 paragraphs ---")
    for i, para in enumerate(doc.paragraphs[:20]):
        print(f"{i}: {para.text[:100]}")
except Exception as e:
    print("Error reading as DOCX directly:", e)
    # Check if there is an alternative library or command we can run to convert/read it
