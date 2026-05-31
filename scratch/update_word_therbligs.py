import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
import os

def set_cell_shading(cell, color_hex):
    shd_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shd_xml))

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin_name, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{margin_name}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table):
    tblPr = table._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top', 'bottom'):
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '12')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '555555')
        tblBorders.append(border)
    insideH = OxmlElement('w:insideH')
    insideH.set(qn('w:val'), 'single')
    insideH.set(qn('w:sz'), '4')
    insideH.set(qn('w:space'), '0')
    insideH.set(qn('w:color'), 'AAAAAA')
    tblBorders.append(insideH)
    for edge in ('left', 'right', 'insideV'):
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'), 'none')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def add_styled_paragraph(target_para, text, bold=False, italic=False, space_before=0, space_after=6, font_size=11):
    new_p = target_para.insert_paragraph_before()
    new_p.paragraph_format.space_before = Pt(space_before)
    new_p.paragraph_format.space_after = Pt(space_after)
    new_p.paragraph_format.line_spacing = 1.15
    new_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    run = new_p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic
    return new_p

def add_styled_subheading(target_para, text):
    p = add_styled_paragraph(target_para, text, bold=True, space_before=12, space_after=4, font_size=12)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def main():
    file_path = r"c:\Users\dsant\Desktop\U CATOLICA\ING INDUSTRIAL\SEMESTRES INDUSTRIAL\segundo semestre\ingenieria de metodos\Corte 1\ingenieria de metodos\Informe Final Métodos.docx"
    
    if not os.path.exists(file_path):
        print("Error: Documento no existe.")
        return
        
    doc = docx.Document(file_path)
    print("Document loaded.")

    # 1. Encontrar el párrafo para insertar la sección 1.2 (antes del Procedimiento 2)
    p2_idx = -1
    for idx, p in enumerate(doc.paragraphs):
        txt = p.text.strip().lower()
        if "2. procedimiento grulla" in txt:
            p2_idx = idx
            break
            
    if p2_idx == -1:
        print("Error: No se encontro el parrafo '2. Procedimiento Grulla Tipo I (Basica)'")
        return

    p2_para = doc.paragraphs[p2_idx]
    
    print(f"Insertando Seccion 1.2 antes de: '{p2_para.text}'")
    
    add_styled_subheading(p2_para, "1.2. Analisis de Micro-movimientos (Therbligs)")
    add_styled_paragraph(p2_para,
        "El analisis de micro-movimientos (Therbligs), formulado originalmente por Frank y Lillian Gilbreth, "
        "descompone la actividad manual en elementos basicos que describen las acciones fisicas y cognitivas del operario. "
        "Mediante el pipeline de vision artificial de CronoGrulla (MediaPipe Holistic) y el clasificador SVM en tiempo real, "
        "se logro catalogar de forma objetiva la secuencia de micro-movimientos para cada una de las 10 tareas del primer "
        "procedimiento de plegado. La Tabla 2.1 sintetiza los Therbligs dominantes detectados para este proceso, categorizados "
        "segun su naturaleza en Therbligs eficientes (agregan valor directo al avance de la operacion)."
    )

    # Crear Tabla 2.1 de Therbligs de Grulla Básica
    t2_1 = doc.add_table(rows=11, cols=4)
    t2_1.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t2_1)
    
    # Mover XML
    parent = p2_para._p.getparent()
    parent.insert(parent.index(p2_para._p), t2_1._tbl)
    
    headers = ["Paso / Elemento", "Actividad del Metodo", "Therblig Predominante", "Tipo de Therblig"]
    for idx, text in enumerate(headers):
        cell = t2_1.cell(0, idx)
        cell.text = text
        set_cell_shading(cell, "F2F4F7")
        set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.name = "Arial"
        run.font.bold = True
        run.font.size = Pt(9.5)

    data_t2_1 = [
        ["Paso 1", "Posicion inicial", "🖐️ SOLTAR (Release - RL)", "Eficiente"],
        ["Paso 2", "Doblar en diagonal", "✊ TOMAR (Grasp - G)", "Eficiente"],
        ["Paso 3", "Segundo doblez", "✊ TOMAR (Grasp - G)", "Eficiente"],
        ["Paso 4", "Formar la base", "✊ TOMAR (Grasp - G)", "Eficiente"],
        ["Paso 5", "Repetir del otro lado", "✊ TOMAR (Grasp - G)", "Eficiente"],
        ["Paso 6", "Formar una especie de pentagono", "✊ TOMAR (Grasp - G)", "Eficiente"],
        ["Paso 7", "Formar el cuerpo", "✊ TOMAR (Grasp - G)", "Eficiente"],
        ["Paso 8", "Formar el cuello", "✊ TOMAR (Grasp - G)", "Eficiente"],
        ["Paso 9", "Formar la cabeza", "🖐️ SOLTAR (Release - RL)", "Eficiente"],
        ["Paso 10", "Ajustar la figura", "🖐️ SOLTAR (Release - RL)", "Eficiente"]
    ]

    for row_idx, row_data in enumerate(data_t2_1):
        for col_idx, text in enumerate(row_data):
            cell = t2_1.cell(row_idx + 1, col_idx)
            cell.text = text
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            if row_idx % 2 == 1:
                set_cell_shading(cell, "FAFAFA")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if col_idx >= 1 else WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.font.name = "Arial"
            run.font.size = Pt(9)

    p_spacer1 = p2_para.insert_paragraph_before()
    p_spacer1.paragraph_format.space_before = Pt(4)
    p_spacer1.paragraph_format.space_after = Pt(8)

    # 2. Encontrar el párrafo para insertar la sección 2.2 (antes de los Informes del INSST)
    insst_idx = -1
    for idx, p in enumerate(doc.paragraphs):
        txt = p.text.strip().lower()
        if "3. informes del insst" in txt:
            insst_idx = idx
            break
            
    if insst_idx == -1:
        print("Error: No se encontro el parrafo '3. Informes del INSST'")
        return

    insst_para = doc.paragraphs[insst_idx]
    
    print(f"Insertando Seccion 2.2 antes de: '{insst_para.text}'")
    
    add_styled_subheading(insst_para, "2.2. Analisis de Micro-movimientos (Therbligs)")
    add_styled_paragraph(insst_para,
        "En la variante intermedia de plegado de la grulla (que consta de 12 pasos y demanda una destreza motora fina "
        "considerablemente mayor), se aplico el mismo pipeline neuronal de CronoGrulla. La Tabla 2.2 expone la secuencia "
        "de Therbligs eficaces evaluados en tiempo real por el clasificador. Este analisis cuantitativo de micro-movimientos "
        "evidencia que la mayor parte del ciclo se compone de Therbligs eficientes de manipulacion y plegado, requiriendo "
        "una coordinacion bimanual extrema y alta concentracion en los cuellos de botella del ensamble."
    )

    # Crear Tabla 2.2 de Therbligs de Grulla Intermedia
    t2_2 = doc.add_table(rows=13, cols=4)
    t2_2.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t2_2)
    
    # Mover XML
    parent2 = insst_para._p.getparent()
    parent2.insert(parent2.index(insst_para._p), t2_2._tbl)
    
    for idx, text in enumerate(headers):
        cell = t2_2.cell(0, idx)
        cell.text = text
        set_cell_shading(cell, "F2F4F7")
        set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.name = "Arial"
        run.font.bold = True
        run.font.size = Pt(9.5)

    data_t2_2 = [
        ["Paso 1", "Diagonal y mitad", "✊ TOMAR (Grasp - G)", "Eficiente"],
        ["Paso 2", "Pliegues cruzados", "✊ TOMAR (Grasp - G)", "Eficiente"],
        ["Paso 3", "Juntar esquinas (Base)", "✊ TOMAR (Grasp - G)", "Eficiente"],
        ["Paso 4", "Marcar solapas", "✊ TOMAR (Grasp - G)", "Eficiente"],
        ["Paso 5", "Marcar punta superior", "✊ TOMAR (Grasp - G)", "Eficiente"],
        ["Paso 6", "Abrir solapa superior", "🖐️ SOLTAR (Release - RL)", "Eficiente"],
        ["Paso 7", "Repetir cara posterior", "✊ TOMAR (Grasp - G)", "Eficiente"],
        ["Paso 8", "Solapas al centro", "✊ TOMAR (Grasp - G)", "Eficiente"],
        ["Paso 9", "Repetir lado opuesto", "✊ TOMAR (Grasp - G)", "Eficiente"],
        ["Paso 10", "Marcar patas inf.", "✊ TOMAR (Grasp - G)", "Eficiente"],
        ["Paso 11", "Pliegue invertido", "✊ TOMAR (Grasp - G)", "Eficiente"],
        ["Paso 12", "Cabeza y alas", "🖐️ SOLTAR (Release - RL)", "Eficiente"]
    ]

    for row_idx, row_data in enumerate(data_t2_2):
        for col_idx, text in enumerate(row_data):
            cell = t2_2.cell(row_idx + 1, col_idx)
            cell.text = text
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            if row_idx % 2 == 1:
                set_cell_shading(cell, "FAFAFA")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if col_idx >= 1 else WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.font.name = "Arial"
            run.font.size = Pt(9)

    p_spacer2 = insst_para.insert_paragraph_before()
    p_spacer2.paragraph_format.space_before = Pt(4)
    p_spacer2.paragraph_format.space_after = Pt(8)

    doc.save(file_path)
    print("Success! Word document modified.")

if __name__ == "__main__":
    main()
