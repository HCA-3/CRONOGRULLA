import docx

file_path = r"c:\Users\dsant\Desktop\U CATOLICA\ING INDUSTRIAL\SEMESTRES INDUSTRIAL\segundo semestre\ingenieria de metodos\Corte 1\ingenieria de metodos\Informe Final Métodos.docx"
doc = docx.Document(file_path)

print("Paragraphs between 44 and 65:")
for i in range(44, min(65, len(doc.paragraphs))):
    print(f"{i}: {doc.paragraphs[i].text}")
