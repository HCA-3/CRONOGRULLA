import os
import docx

def main():
    p = r"C:\Users\dsant\Desktop\esta\Articulo_CronoGrulla_IEEE.docx"
    if not os.path.exists(p):
        print("El archivo no existe")
        return
        
    doc = docx.Document(p)
    print("Abriendo archivo Word...")

    def decode_corrupt_chars(text):
        if not text:
            return ""
        result = []
        for c in text:
            code = ord(c)
            # Mapeo manual estricto de caracteres según sus valores CP1252 / ISO-8859-1 decodificados erróneamente por Word
            if code == 225:   # á
                result.append('á')
            elif code == 233: # é
                result.append('é')
            elif code == 237: # í
                result.append('í')
            elif code == 243: # ó
                result.append('ó')
            elif code == 250: # ú
                result.append('ú')
            elif code == 241: # ñ
                result.append('ñ')
            elif code == 193: # Á
                result.append('Á')
            elif code == 201: # É
                result.append('É')
            elif code == 205: # Í
                result.append('Í')
            elif code == 211: # Ó
                result.append('Ó')
            elif code == 218: # Ú
                result.append('Ú')
            elif code == 209: # Ñ
                result.append('Ñ')
            elif code == 173: # Guion suave, eliminar
                continue
            elif code == 150: # Guion largo
                result.append('—')
            else:
                result.append(c)
        return "".join(result)

    # Recorrer todos los párrafos y aplicar correcciones de codificación
    for paragraph in doc.paragraphs:
        txt = paragraph.text
        txt = decode_corrupt_chars(txt)
        
        # Correcciones ortográficas generales y de formato
        txt = txt.replace("estudio", "estudio").replace("estudios", "estudios")
        txt = txt.replace("estaciones", "estaciones").replace("estación", "estación")
        txt = txt.replace("investigaciones", "investigaciones")
        txt = txt.replace("ptimo", "óptimo").replace("ptima", "óptima")
        txt = txt.replace("ptimos", "óptimos").replace("ptimas", "óptimas")
        txt = txt.replace("ptimal", "optimal")
        paragraph.text = txt

    # Recorrer todas las tablas y aplicar correcciones de codificación
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                txt = cell.text
                txt = decode_corrupt_chars(txt)
                txt = txt.replace("estudio", "estudio").replace("estudios", "estudios")
                txt = txt.replace("ptimo", "óptimo").replace("ptima", "óptima")
                cell.text = txt

    doc.save(p)
    print("Guardado final completado con tildes perfectas decodificadas.")

if __name__ == "__main__":
    main()
